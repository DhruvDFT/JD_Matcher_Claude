#!/usr/bin/env python3
"""
JD Profile Matcher - Complete Single File Application
A web application that matches candidate profiles against job descriptions.

Setup Instructions:
1. Save this file as 'app.py'
2. Create requirements.txt with the dependencies listed below
3. Run: pip install -r requirements.txt
4. Run: python app.py
5. Deploy to Railway/Heroku by pushing to GitHub

Required requirements.txt:
Flask==2.3.3
PyPDF2==3.0.1
python-docx==0.8.11
nltk==3.8.1
scikit-learn==1.3.0
numpy==1.24.3
werkzeug==2.3.7
gunicorn==21.2.0
click==8.1.7
"""

import os
import re
import json
import ssl
from typing import Dict, List, Set
from werkzeug.utils import secure_filename
from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify

# External libraries with error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

# NLTK setup with error handling
def setup_nltk():
    """Setup NLTK data with proper error handling for cloud deployment"""
    try:
        import nltk
        # Handle SSL issues in cloud environments
        try:
            _create_unverified_https_context = ssl._create_unverified_context
        except AttributeError:
            pass
        else:
            ssl._create_default_https_context = _create_unverified_https_context
        
        # Download required NLTK data
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        return True
    except Exception as e:
        print(f"NLTK setup warning: {e}")
        return False

NLTK_AVAILABLE = setup_nltk()

if NLTK_AVAILABLE:
    try:
        from nltk.corpus import stopwords
        from nltk.tokenize import word_tokenize
    except ImportError:
        NLTK_AVAILABLE = False

# Flask app setup
app = Flask(__name__)
app.secret_key = 'your-secret-key-change-for-production'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Text processing functions
def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats with detailed logging"""
    _, ext = os.path.splitext(file_path.lower())
    print(f"DEBUG: Extracting text from {ext} file: {file_path}")
    
    if ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                print(f"DEBUG: TXT file read successfully, length: {len(content)}")
                return content
        except Exception as e:
            print(f"DEBUG: Error reading TXT file: {e}")
            raise ValueError(f"Could not read TXT file: {e}")
    
    elif ext == '.pdf':
        if not PDF_AVAILABLE:
            raise ValueError("PDF support not available. Please install PyPDF2 or use TXT files.")
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"DEBUG: PDF has {len(pdf_reader.pages)} pages")
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"DEBUG: Page {i+1} extracted {len(page_text)} characters")
                print(f"DEBUG: PDF extraction complete, total length: {len(text)}")
                return text
        except Exception as e:
            print(f"DEBUG: Error reading PDF file: {e}")
            raise ValueError(f"Could not read PDF file: {e}")
    
    elif ext in ['.doc', '.docx']:
        if not DOCX_AVAILABLE:
            raise ValueError("DOC/DOCX support not available. Please install python-docx or use TXT files.")
        
        try:
            doc = docx.Document(file_path)
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
            print(f"DEBUG: DOC/DOCX extraction complete, {len(doc.paragraphs)} paragraphs, total length: {len(text)}")
            return text
        except Exception as e:
            print(f"DEBUG: Error reading DOC/DOCX file: {e}")
            raise ValueError(f"Could not read DOC/DOCX file: {e}")
    
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .txt, .pdf, .doc, .docx")

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\-\+\#\.]', ' ', text)
    return text.strip()

def extract_skills(text: str) -> Set[str]:
    """Extract skills from text using comprehensive patterns"""
    text = text.lower()
    
    # Comprehensive skill patterns including VLSI and Embedded
    skill_patterns = [
        # Programming Languages
        r'\b(?:python|java|javascript|c\+\+|c#|php|ruby|go|rust|swift|kotlin|c|verilog|vhdl|systemverilog|assembly)\b',
        
        # Web Technologies
        r'\b(?:react|angular|vue|node\.?js|express|django|flask|spring)\b',
        
        # Databases
        r'\b(?:sql|mysql|postgresql|mongodb|redis|elasticsearch)\b',
        
        # Cloud & DevOps
        r'\b(?:aws|azure|gcp|docker|kubernetes|jenkins|git)\b',
        
        # AI/ML
        r'\b(?:machine learning|deep learning|ai|data science|nlp)\b',
        
        # Frontend
        r'\b(?:html|css|bootstrap|tailwind|scss|sass)\b',
        
        # Methodologies
        r'\b(?:agile|scrum|kanban|devops|ci/cd)\b',
        
        # VLSI & IC Design
        r'\b(?:vlsi|asic|fpga|rtl|synthesis|place and route|dft|sta|timing analysis)\b',
        r'\b(?:soc|system on chip|ic design|analog design|digital design)\b',
        r'\b(?:verification|uvm|testbench|assertion|coverage|simulation)\b',
        r'\b(?:spice|hspice|spectre|virtuoso|calibre|primetime|dc compiler)\b',
        r'\b(?:amba|axi|ahb|apb|wishbone|avalon)\b',
        
        # Embedded Systems
        r'\b(?:embedded|microcontroller|mcu|arm|risc-v|cortex|atmega|pic|esp32|arduino)\b',
        r'\b(?:firmware|bootloader|rtos|freertos|embedded c|bare metal)\b',
        r'\b(?:pcb design|schematic|altium|kicad|eagle|cadence|mentor graphics)\b',
        r'\b(?:uart|spi|i2c|can|usb|ethernet|bluetooth|wifi|zigbee)\b',
        r'\b(?:adc|dac|pwm|gpio|timer|interrupt|dma)\b',
        
        # Hardware Protocols & Standards
        r'\b(?:pcie|ddr|usb 3\.0|mipi|hdmi|displayport)\b',
        
        # Tools & Software
        r'\b(?:matlab|simulink|labview|quartus|vivado|ise|modelsim|questasim)\b',
        r'\b(?:synopsys|cadence|mentor|xilinx|altera|intel)\b'
    ]
    
    skills = set()
    for pattern in skill_patterns:
        matches = re.findall(pattern, text)
        skills.update(matches)
    
    # Extract skills from common sections
    skill_sections = re.findall(r'(?:skills?|technologies?|tools?|expertise|experience)[:\-\s]*(.*?)(?:\n\n|\n[A-Z]|$)', text, re.IGNORECASE | re.DOTALL)
    for section in skill_sections:
        section_skills = re.split(r'[,\n\|•\-\*\(\)\/]', section)
        for skill in section_skills:
            skill = skill.strip()
            if skill and len(skill) > 1 and len(skill) < 50:
                skills.add(skill.lower())
    
    # Additional VLSI/Embedded specific extractions
    vlsi_embedded_terms = [
        'vlsi', 'asic', 'fpga', 'rtl design', 'verilog', 'vhdl', 'systemverilog',
        'embedded systems', 'microcontroller', 'firmware', 'rtos', 'embedded c',
        'arm cortex', 'risc-v', 'soc design', 'ic design', 'analog design',
        'digital design', 'pcb design', 'schematic design', 'verification',
        'dft', 'sta', 'synthesis', 'place and route', 'timing analysis',
        'uart', 'spi', 'i2c', 'can bus', 'usb', 'ethernet', 'bluetooth',
        'adc', 'dac', 'pwm', 'gpio', 'interrupt handling', 'bootloader',
        'bare metal programming', 'low power design', 'signal integrity',
        'emi/emc', 'cadence', 'synopsys', 'mentor graphics', 'xilinx',
        'altera', 'quartus', 'vivado', 'modelsim', 'questasim'
    ]
    
    for term in vlsi_embedded_terms:
        if term in text:
            skills.add(term)
    
    return skills

def extract_experience_years(text: str) -> int:
    """Extract years of experience from text"""
    text = text.lower()
    
    patterns = [
        r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?experience',
        r'(\d+)\s*\+?\s*yrs?\s+(?:of\s+)?experience',
        r'experience\s*:?\s*(\d+)\s*\+?\s*years?',
        r'(\d+)\s*\+?\s*years?\s+in\s+',
    ]
    
    years = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        years.extend([int(match) for match in matches])
    
    return max(years) if years else 0

def extract_education(text: str) -> List[str]:
    """Extract education information"""
    text = text.lower()
    
    degrees = []
    degree_patterns = [
        r'\b(?:bachelor|b\.?[sca]\.?|bs|ba|bsc|bca|be|btech)\b',
        r'\b(?:master|m\.?[sca]\.?|ms|ma|msc|mca|me|mtech|mba)\b',
        r'\b(?:phd|ph\.?d\.?|doctorate|doctoral)\b',
        r'\b(?:diploma|certificate)\b'
    ]
    
    for pattern in degree_patterns:
        if re.search(pattern, text):
            degrees.append(pattern)
    
    return degrees

def extract_keywords(text: str) -> Set[str]:
    """Extract important keywords with fallback for cloud deployment"""
    if NLTK_AVAILABLE:
        try:
            stop_words = set(stopwords.words('english'))
            word_tokens = word_tokenize(text.lower())
            
            keywords = {word for word in word_tokens 
                       if word not in stop_words 
                       and len(word) > 2 
                       and word.isalpha()}
            
            return keywords
        except Exception as e:
            print(f"NLTK processing failed: {e}")
    
    # Fallback method without NLTK
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    simple_stopwords = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'been', 'have', 'has', 'had', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those'}
    return {word for word in words if word not in simple_stopwords}

def parse_jd(text: str) -> Dict:
    """Parse job description and extract key information"""
    text = clean_text(text)
    
    return {
        'raw_text': text,
        'required_skills': extract_skills(text),
        'min_experience': extract_experience_years(text),
        'education_requirements': extract_education(text),
        'keywords': extract_keywords(text)
    }

def parse_profile(text: str) -> Dict:
    """Parse candidate profile and extract key information"""
    text = clean_text(text)
    
    return {
        'raw_text': text,
        'skills': extract_skills(text),
        'experience_years': extract_experience_years(text),
        'education': extract_education(text),
        'keywords': extract_keywords(text)
    }

# Matching functions
def calculate_match_score(jd_data: Dict, profile_data: Dict) -> float:
    """Calculate match score between JD and profile"""
    
    weights = {
        'skills': 0.4,
        'experience': 0.2,
        'education': 0.1,
        'text_similarity': 0.3
    }
    
    scores = {}
    
    # 1. Skills matching
    jd_skills = jd_data.get('required_skills', set())
    profile_skills = profile_data.get('skills', set())
    
    if jd_skills:
        skill_overlap = len(jd_skills.intersection(profile_skills))
        scores['skills'] = skill_overlap / len(jd_skills)
    else:
        scores['skills'] = 0.5
    
    # 2. Experience matching
    min_exp = jd_data.get('min_experience', 0)
    candidate_exp = profile_data.get('experience_years', 0)
    
    if min_exp == 0:
        scores['experience'] = 1.0
    elif candidate_exp >= min_exp:
        scores['experience'] = 1.0
    else:
        scores['experience'] = candidate_exp / min_exp
    
    # 3. Education matching
    jd_education = set(jd_data.get('education_requirements', []))
    profile_education = set(profile_data.get('education', []))
    
    if jd_education:
        if jd_education.intersection(profile_education):
            scores['education'] = 1.0
        else:
            scores['education'] = 0.3
    else:
        scores['education'] = 1.0
    
    # 4. Text similarity using TF-IDF (with fallback)
    jd_text = jd_data.get('raw_text', '')
    profile_text = profile_data.get('raw_text', '')
    
    if jd_text and profile_text and SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
            tfidf_matrix = vectorizer.fit_transform([jd_text, profile_text])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
            scores['text_similarity'] = similarity[0][0]
        except:
            scores['text_similarity'] = 0.5
    else:
        # Fallback: simple keyword overlap
        jd_keywords = jd_data.get('keywords', set())
        profile_keywords = profile_data.get('keywords', set())
        if jd_keywords:
            overlap = len(jd_keywords.intersection(profile_keywords))
            scores['text_similarity'] = min(overlap / len(jd_keywords), 1.0)
        else:
            scores['text_similarity'] = 0.5
    
    # Calculate weighted average
    total_score = sum(scores[key] * weights[key] for key in weights)
    return min(total_score, 1.0)

def rank_profiles(matches: List[Dict]) -> List[Dict]:
    """Rank profiles by match score"""
    return sorted(matches, key=lambda x: x['score'], reverse=True)

# HTML Templates
INDEX_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>JD Profile Matcher</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; background-color: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; padding: 20px; }
        header { text-align: center; margin-bottom: 30px; }
        header h1 { color: #2c3e50; margin-bottom: 10px; }
        header p { color: #666; font-size: 1.1em; }
        .upload-form { background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); margin-bottom: 30px; }
        .form-group { margin-bottom: 20px; }
        .form-group label { display: block; margin-bottom: 5px; font-weight: 600; color: #2c3e50; }
        .form-group input[type="file"] { width: 100%; padding: 10px; border: 2px dashed #ddd; border-radius: 5px; background: #f9f9f9; }
        .form-group small { color: #666; font-size: 0.9em; }
        .btn-primary { background: #3498db; color: white; padding: 12px 30px; border: none; border-radius: 5px; font-size: 1.1em; cursor: pointer; transition: background 0.3s; }
        .btn-primary:hover { background: #2980b9; }
        .btn-secondary { background: #95a5a6; color: white; padding: 8px 20px; text-decoration: none; border-radius: 5px; font-size: 0.9em; transition: background 0.3s; }
        .btn-secondary:hover { background: #7f8c8d; }
        .messages .alert { padding: 10px; margin-bottom: 20px; border-radius: 5px; }
        .alert-error { background: #e74c3c; color: white; }
        .alert-success { background: #27ae60; color: white; }
        .api-info { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); margin-top: 30px; }
        .api-info code { background: #f1f2f6; padding: 2px 5px; border-radius: 3px; font-family: 'Courier New', monospace; }
        .api-info pre { background: #2c3e50; color: white; padding: 15px; border-radius: 5px; overflow-x: auto; margin-top: 10px; }
        @media (max-width: 768px) { .container { padding: 10px; } }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>JD Profile Matcher</h1>
            <p>Upload job description and candidate profiles to find the best matches</p>
        </header>

        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                <div class="messages">
                    {% for category, message in messages %}
                        <div class="alert alert-{{ category }}">{{ message }}</div>
                    {% endfor %}
                </div>
            {% endif %}
        {% endwith %}

        <form method="POST" action="{{ url_for('upload_files') }}" enctype="multipart/form-data" class="upload-form">
            <div class="form-group">
                <label for="jd_file">Job Description (PDF, DOC, DOCX, TXT)</label>
                <input type="file" id="jd_file" name="jd_file" accept=".pdf,.doc,.docx,.txt" required>
                <small>Upload the job description file</small>
            </div>

            <div class="form-group">
                <label for="profile_files">Candidate Profiles (PDF, DOC, DOCX, TXT)</label>
                <input type="file" id="profile_files" name="profile_files" accept=".pdf,.doc,.docx,.txt" multiple required>
                <small>Select multiple candidate profile files (hold Ctrl/Cmd to select multiple)</small>
            </div>

            <button type="submit" class="btn-primary">🚀 Start Matching</button>
            
            <div id="loading" style="display: none; text-align: center; margin-top: 20px;">
                <p>🔄 Processing files and matching profiles...</p>
                <p>This may take a few moments for large files.</p>
            </div>
        </form>

        <script>
            document.querySelector('form').addEventListener('submit', function() {
                document.getElementById('loading').style.display = 'block';
                document.querySelector('.btn-primary').disabled = true;
                document.querySelector('.btn-primary').textContent = 'Processing...';
            });
        </script>

        <div class="api-info">
            <h3>🚀 Deployment Ready</h3>
            <p><strong>VLSI & Embedded Skills:</strong> ✅ Comprehensive support for VLSI, ASIC, FPGA, RTL, Embedded Systems, Firmware, and Hardware Design skills</p>
            <p><strong>Railway Deployment:</strong> ✅ Optimized for cloud deployment with robust error handling</p>
            
            <h3>API Usage</h3>
            <p>You can also use the API endpoint:</p>
            <code>POST /api/match</code>
            <pre>{
  "jd_text": "Job description text...",
  "profiles": ["Profile 1 text...", "Profile 2 text..."]
}</pre>
        </div>
    </div>
</body>
</html>
"""

RESULTS_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Matching Results - JD Profile Matcher</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; background-color: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; padding: 20px; }
        header { text-align: center; margin-bottom: 30px; }
        header h1 { color: #2c3e50; margin-bottom: 10px; }
        .btn-secondary { background: #95a5a6; color: white; padding: 8px 20px; text-decoration: none; border-radius: 5px; font-size: 0.9em; transition: background 0.3s; }
        .btn-secondary:hover { background: #7f8c8d; }
        .jd-summary { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); margin-bottom: 30px; }
        .summary-grid { display: grid; gap: 15px; }
        .summary-item { padding: 10px 0; }
        .skills-list { display: flex; flex-wrap: wrap; gap: 5px; margin-top: 5px; }
        .skill-tag { background: #ecf0f1; padding: 4px 8px; border-radius: 4px; font-size: 0.9em; border: 1px solid #bdc3c7; }
        .skill-tag.matched { background: #d5f4e6; border-color: #27ae60; color: #27ae60; }
        .results-section { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .profile-card { border: 1px solid #ddd; border-radius: 8px; margin-bottom: 20px; padding: 20px; background: #fafafa; }
        .profile-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 15px; }
        .profile-header h3 { color: #2c3e50; margin: 0; }
        .match-score { text-align: center; }
        .score-circle { width: 60px; height: 60px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; color: #2c3e50; position: relative; }
        .score-circle::before { content: ''; position: absolute; width: 40px; height: 40px; background: white; border-radius: 50%; z-index: -1; }
        .detail-row { margin-bottom: 10px; display: flex; flex-wrap: wrap; gap: 10px; }
        .detail-row strong { min-width: 100px; color: #2c3e50; }
        @media (max-width: 768px) { .container { padding: 10px; } .profile-header { flex-direction: column; align-items: flex-start; gap: 10px; } .detail-row { flex-direction: column; gap: 5px; } .detail-row strong { min-width: auto; } }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Matching Results</h1>
            <a href="{{ url_for('index') }}" class="btn-secondary">New Search</a>
        </header>

        <div class="jd-summary">
            <h2>Job Requirements Summary</h2>
            <div class="summary-grid">
                <div class="summary-item">
                    <strong>Required Skills:</strong>
                    {% if jd_data.required_skills %}
                        <span class="skills-list">
                            {% for skill in jd_data.required_skills %}
                                <span class="skill-tag">{{ skill }}</span>
                            {% endfor %}
                        </span>
                    {% else %}
                        <span>No specific skills identified</span>
                    {% endif %}
                </div>
                <div class="summary-item">
                    <strong>Minimum Experience:</strong>
                    <span>{{ jd_data.min_experience }} years</span>
                </div>
            </div>
        </div>

        <div class="results-section">
            <h2>Candidate Matches ({{ matches|length }} profiles)</h2>
            
            {% for match in matches %}
                <div class="profile-card">
                    <div class="profile-header">
                        <h3>{{ match.profile.filename }}</h3>
                        <div class="match-score">
                            <div class="score-circle" style="background: conic-gradient(#4CAF50 {{ match.percentage }}%, #ddd {{ match.percentage }}%);">
                                <span>{{ match.percentage }}%</span>
                            </div>
                        </div>
                    </div>
                    
                    <div class="profile-details">
                        <div class="detail-row">
                            <strong>Experience:</strong>
                            <span>{{ match.profile.experience_years }} years</span>
                        </div>
                        
                        <div class="detail-row">
                            <strong>Skills:</strong>
                            <div class="skills-list">
                                {% for skill in match.profile.skills %}
                                    <span class="skill-tag {% if skill in jd_data.required_skills %}matched{% endif %}">
                                        {{ skill }}
                                    </span>
                                {% endfor %}
                            </div>
                        </div>
                        
                        {% if match.profile.education %}
                        <div class="detail-row">
                            <strong>Education:</strong>
                            <span>{{ match.profile.education|join(', ') }}</span>
                        </div>
                        {% endif %}
                    </div>
                </div>
            {% endfor %}
        </div>
    </div>
</body>
</html>
"""

# Flask routes
@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

@app.route('/upload', methods=['POST'])
def upload_files():
    try:
        print("=== DEBUG: Starting file upload process ===")
        
        # Check if JD file is uploaded
        if 'jd_file' not in request.files:
            flash('No JD file uploaded', 'error')
            return redirect(url_for('index'))
        
        jd_file = request.files['jd_file']
        if jd_file.filename == '':
            flash('No JD file selected', 'error')
            return redirect(url_for('index'))
        
        print(f"DEBUG: JD file received: {jd_file.filename}")
        
        # Check if profile files are uploaded
        if 'profile_files' not in request.files:
            flash('No profile files uploaded', 'error')
            return redirect(url_for('index'))
        
        profile_files = request.files.getlist('profile_files')
        if not profile_files or all(f.filename == '' for f in profile_files):
            flash('No profile files selected', 'error')
            return redirect(url_for('index'))
        
        print(f"DEBUG: Profile files received: {[f.filename for f in profile_files]}")
        
        # Process JD file
        if jd_file and allowed_file(jd_file.filename):
            jd_filename = secure_filename(jd_file.filename)
            jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
            jd_file.save(jd_path)
            
            print(f"DEBUG: Processing JD file: {jd_path}")
            
            # Extract and parse JD
            try:
                jd_text = extract_text_from_file(jd_path)
                print(f"DEBUG: JD text extracted, length: {len(jd_text)}")
                print(f"DEBUG: JD text preview: {jd_text[:200]}...")
                
                jd_data = parse_jd(jd_text)
                print(f"DEBUG: JD parsed - Skills: {len(jd_data.get('required_skills', set()))}, Experience: {jd_data.get('min_experience', 0)}")
                
            except Exception as e:
                print(f"DEBUG: Error processing JD file: {e}")
                flash(f'Error reading JD file: {str(e)}', 'error')
                return redirect(url_for('index'))
            
            # Process profile files
            profiles = []
            for i, profile_file in enumerate(profile_files):
                if profile_file and allowed_file(profile_file.filename):
                    profile_filename = secure_filename(profile_file.filename)
                    profile_path = os.path.join(app.config['UPLOAD_FOLDER'], profile_filename)
                    profile_file.save(profile_path)
                    
                    print(f"DEBUG: Processing profile {i+1}: {profile_filename}")
                    
                    try:
                        # Extract and parse profile
                        profile_text = extract_text_from_file(profile_path)
                        print(f"DEBUG: Profile {i+1} text extracted, length: {len(profile_text)}")
                        
                        profile_data = parse_profile(profile_text)
                        profile_data['filename'] = profile_filename
                        profiles.append(profile_data)
                        
                        print(f"DEBUG: Profile {i+1} parsed - Skills: {len(profile_data.get('skills', set()))}, Experience: {profile_data.get('experience_years', 0)}")
                        
                    except Exception as e:
                        print(f"DEBUG: Error processing profile {profile_filename}: {e}")
                        flash(f'Error reading profile {profile_filename}: {str(e)}', 'error')
                    finally:
                        # Clean up uploaded file
                        if os.path.exists(profile_path):
                            os.remove(profile_path)
            
            # Clean up JD file
            if os.path.exists(jd_path):
                os.remove(jd_path)
            
            print(f"DEBUG: Successfully processed {len(profiles)} profiles")
            
            
            if not profiles:
                flash('No valid profiles could be processed. Please check your file formats.', 'error')
                return redirect(url_for('index'))
            
            # Calculate matches
            print("DEBUG: Starting matching process...")
            matches = []
            for i, profile in enumerate(profiles):
                try:
                    score = calculate_match_score(jd_data, profile)
                    match_data = {
                        'profile': profile,
                        'score': score,
                        'percentage': round(score * 100, 1)
                    }
                    matches.append(match_data)
                    print(f"DEBUG: Profile {i+1} ({profile.get('filename', 'Unknown')}) - Score: {score:.3f} ({match_data['percentage']}%)")
                except Exception as e:
                    print(f"DEBUG: Error calculating match for profile {i+1}: {e}")
                    # Add with 0 score if calculation fails
                    matches.append({
                        'profile': profile,
                        'score': 0.0,
                        'percentage': 0.0
                    })
            
            # Rank profiles by match score
            ranked_matches = rank_profiles(matches)
            print(f"DEBUG: Ranked {len(ranked_matches)} matches")
            
            # Debug output for template
            print("DEBUG: Preparing template data...")
            print(f"DEBUG: JD Skills: {list(jd_data.get('required_skills', set()))[:5]}...")  # Show first 5 skills
            print(f"DEBUG: Top match: {ranked_matches[0]['percentage']}%" if ranked_matches else "No matches")
            
            return render_template_string(RESULTS_HTML, 
                                        jd_data=jd_data, 
                                        matches=ranked_matches)
        
        else:
            flash('Invalid file type. Please upload TXT, PDF, DOC, or DOCX files.', 'error')
            return redirect(url_for('index'))
            
    except Exception as e:
        print(f"DEBUG: Fatal error in upload_files: {e}")
        import traceback
        traceback.print_exc()
        flash(f'Error processing files: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/api/match', methods=['POST'])
def api_match():
    """API endpoint for programmatic access"""
    try:
        data = request.get_json()
        jd_text = data.get('jd_text', '')
        profiles_text = data.get('profiles', [])
        
        if not jd_text or not profiles_text:
            return jsonify({'error': 'Missing JD text or profiles'}), 400
        
        # Parse JD
        jd_data = parse_jd(jd_text)
        
        # Process profiles
        matches = []
        for i, profile_text in enumerate(profiles_text):
            profile_data = parse_profile(profile_text)
            profile_data['id'] = i
            score = calculate_match_score(jd_data, profile_data)
            matches.append({
                'profile_id': i,
                'score': score,
                'percentage': round(score * 100, 1),
                'profile_data': profile_data
            })
        
        # Rank profiles
        ranked_matches = rank_profiles(matches)
        
        return jsonify({
            'jd_data': jd_data,
            'matches': ranked_matches
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)

"""
DEPLOYMENT INSTRUCTIONS:

1. SAVE THIS FILE:
   - Save as 'app.py'

2. CREATE requirements.txt (FIXED VERSION):
Flask==2.3.2
PyPDF2==3.0.1
python-docx==0.8.11
nltk==3.8.1
scikit-learn==1.3.0
numpy==1.24.3
werkzeug==2.3.6
gunicorn==20.1.0

3. CREATE Procfile (REQUIRED for Railway/Heroku):
web: gunicorn app:app --bind 0.0.0.0:$PORT --workers 1 --timeout 120

4. CREATE railway.json (ALTERNATIVE CONFIG):
{
  "build": {
    "builder": "NIXPACKS"
  },
  "deploy": {
    "startCommand": "gunicorn app:app --host 0.0.0.0 --port $PORT",
    "restartPolicyType": "ON_FAILURE"
  }
}

5. CREATE runtime.txt (SPECIFY PYTHON VERSION):
python-3.9.18

6. CREATE .gitignore:
__pycache__/
*.pyc
uploads/
.env
*.log
venv/

7. MINIMAL requirements.txt (IF ABOVE FAILS):
Flask==2.3.2
gunicorn==20.1.0
werkzeug==2.3.6

8. DEPLOYMENT TROUBLESHOOTING:

ISSUE: pip install fails (exit code 2)
SOLUTIONS (try in order):

A) Use minimal requirements first:
Flask==2.3.2
gunicorn==20.1.0

B) Force older Python version in runtime.txt:
python-3.9.18

C) Alternative Railway deployment method:
- Delete railway.json
- Railway will auto-detect Python
- Uses default pip install

D) Use Docker deployment on Railway:
Create Dockerfile:
FROM python:3.9-slim
WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY . .
EXPOSE 5000
CMD ["gunicorn", "app:app", "--bind", "0.0.0.0:5000"]

9. ALTERNATIVE DEPLOYMENT PLATFORMS:

HEROKU (Most Reliable):
heroku create your-app-name
git push heroku main

RENDER (Free, Easy):
- Connect GitHub to render.com
- Build command: pip install -r requirements.txt
- Start command: gunicorn app:app --bind 0.0.0.0:$PORT

RAILWAY ALTERNATIVE SETUP:
- Remove all config files except app.py and requirements.txt
- Let Railway auto-detect everything
- Use minimal requirements.txt

10. STEP-BY-STEP FIX:

Step 1: Try minimal requirements.txt:
Flask==2.3.2
gunicorn==20.1.0

Step 2: If that works, gradually add:
Flask==2.3.2
gunicorn==20.1.0
PyPDF2==3.0.1
python-docx==0.8.11

Step 3: Then add ML libraries:
Flask==2.3.2
gunicorn==20.1.0
PyPDF2==3.0.1
python-docx==0.8.11
nltk==3.8.1
scikit-learn==1.3.0
numpy==1.24.3

11. EMERGENCY DEPLOYMENT (No external libraries):
If all else fails, the app has fallbacks and will work with just Flask!

Basic requirements.txt:
Flask==2.3.2
gunicorn==20.1.0

The app will:
- ✅ Still work for TXT files
- ✅ Use fallback text processing
- ✅ Provide basic matching
- ❌ No PDF/DOC support (use TXT instead)
- ❌ Simpler text analysis

12. LOCAL TESTING FIRST:
python -m venv venv
source venv/bin/activate  # Windows: venv\Scripts\activate
pip install Flask gunicorn
python app.py
# Test at localhost:5000

13. DEPLOYMENT ORDER OF PREFERENCE:
1. Heroku (most reliable)
2. Render (free tier)
3. Railway (if config issues resolved)
4. Local + ngrok for team sharing

The key is starting with minimal dependencies and adding more once basic deployment works!
"""
